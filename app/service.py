"""
Service layer for the Text Summarizer System using BERT.

This module provides functions to clean text, generate word clouds, summarize text using BERT,
validate URLs, save summaries to DOCX files, fetch articles, and process uploaded files.
"""

import re
from wordcloud import WordCloud
import matplotlib
matplotlib.use('Agg')  # Use the Agg backend to suppress the GUI warning
import matplotlib.pyplot as plt
from summarizer import Summarizer
import docx
from urllib.parse import urlparse
import requests
from newspaper import fulltext
import logging
import pathlib
import os
import PyPDF2
from docx import Document
import uuid

# Configure logging for the service module
logging.basicConfig(level=logging.ERROR)

# Update constants with absolute paths
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DOCX_PATH = os.path.join(BASE_DIR, 'static', 'download', 'file.docx')
WORDCLOUD_PATH = os.path.join(BASE_DIR, 'static', 'img', 'wordcloud', 'wordcloud.png')

def clean_text_and_generate_wordcloud(file_content):
    """
    Clean the input text and generate a word cloud image.

    Args:
        file_content (str): The input text to be cleaned.

    Returns:
        str: The cleaned text.
    """
    file_content = re.sub(r'\d', ' ', file_content)
    file_content = re.sub(r'\W', ' ', file_content)
    file_content = re.sub(r'\s+', ' ', file_content)
    file_content = re.sub(r'\[[0-9]*\]', ' ', file_content)

    try:
        # Create all necessary directories
        os.makedirs(os.path.dirname(WORDCLOUD_PATH), exist_ok=True)
        
        # Close any existing plots
        plt.close('all')
        
        wordcloud = WordCloud(max_font_size=100, max_words=100, background_color="white").generate(file_content)
        plt.figure()
        plt.imshow(wordcloud, interpolation='bilinear')
        plt.axis("off")
        plt.savefig(WORDCLOUD_PATH)
        plt.close()  # Close the plot after saving
    except Exception as e:
        logging.error(f"Error generating word cloud: {e}")
    return file_content

def generate_bert_summary(cleaned_text):
    """
    Generate a summary using the BERT model.

    Args:
        cleaned_text (str): The cleaned input text.

    Returns:
        str: The generated summary with proper sentence separation.
    """
    try:
        model = Summarizer()
        result = model(cleaned_text, min_length=50)
        
        # Improved sentence detection and formatting
        sentences = []
        current_sentence = ""
        
        # Split into words and reconstruct sentences
        words = result.split()
        for word in words:
            current_sentence += word + " "
            # Check for sentence endings
            if word.endswith(('.', '!', '?')) or (word[-1].isalnum() and len(current_sentence) > 50):
                current_sentence = current_sentence.strip()
                if not current_sentence[-1] in '.!?':
                    current_sentence += '.'
                if current_sentence[0].islower():
                    current_sentence = current_sentence[0].upper() + current_sentence[1:]
                sentences.append(current_sentence)
                current_sentence = ""
        
        # Handle any remaining text
        if current_sentence:
            if not current_sentence[-1] in '.!?':
                current_sentence += '.'
            sentences.append(current_sentence.strip())
        
        # Join sentences with proper spacing
        return ' '.join(sentences)
    except Exception as e:
        logging.error(f"Error generating summary: {e}")
        return "Error generating summary"

def url_validator(url):
    """
    Validate the given URL.

    Args:
        url (str): The URL to be validated.

    Returns:
        bool: True if the URL is valid, False otherwise.
    """
    try:
        result = urlparse(url)
        return all([result.scheme, result.netloc, result.path])
    except Exception as e:
        logging.error(f"Error validating URL: {e}")
        return False

def save_docx_file(summary_text):
    """
    Save summary and word cloud image to a DOCX file.

    Args:
        summary_text (str): The summary text to be saved.
    """
    try:
        os.makedirs(os.path.dirname(DOCX_PATH), exist_ok=True)
        mydoc = docx.Document()
        mydoc.add_heading("Summary", 0)
        mydoc.add_paragraph(summary_text)
        mydoc.add_picture(WORDCLOUD_PATH, width=docx.shared.Inches(5), height=docx.shared.Inches(6))
        mydoc.save(DOCX_PATH)
    except Exception as e:
        logging.error(f"Error saving DOCX file: {e}")

def fetch_article(url):
    """
    Fetch the article text from the given URL.

    Args:
        url (str): The URL of the article.

    Returns:
        str: The article text.
    """
    try:
        response = requests.get(url)
        return fulltext(response.text)
    except Exception as e:
        logging.error(f"Error fetching article: {e}")
        return "Error fetching article"

def summarize_url(url):
    """
    Summarize the article from the given URL.

    Args:
        url (str): The URL of the article.

    Returns:
        str: The summary of the article.
    """
    try:
        if not url_validator(url):
            return "Invalid URL"
        article_text = fetch_article(url)
        cleaned_text = clean_text_and_generate_wordcloud(article_text)
        summary_text = generate_bert_summary(cleaned_text)
        save_docx_file(summary_text)
        return summary_text
    except Exception as e:
        logging.error(f"Error summarizing URL: {e}")
        return "Error summarizing URL"

def remove_files():
    """
    Remove previously generated word cloud and DOCX files if present.
    """
    wordcloud_file = pathlib.Path(WORDCLOUD_PATH)
    if wordcloud_file.is_file():
        os.remove(wordcloud_file)

    docx_file = pathlib.Path(DOCX_PATH)
    if docx_file.is_file():
        os.remove(docx_file)

def process_file(uploaded_file):
    """
    Process the uploaded file and return the extracted text.

    Args:
        uploaded_file: The uploaded file.

    Returns:
        str: The extracted text.

    Raises:
        ValueError: If the file type is unsupported.
    """
    suffix = pathlib.Path(uploaded_file.filename).suffix.lower()
    if suffix == ".pdf":
        with open(uploaded_file.filename, 'rb') as pdf:
            pdfReader = PyPDF2.PdfReader(pdf)
            return "".join([page.extract_text() for page in pdfReader.pages if page.extract_text()])
    elif suffix == ".docx":
        doc = docx.Document(uploaded_file.filename)
        return "".join([para.text for para in doc.paragraphs])
    elif suffix == ".txt":
        with open(uploaded_file.filename, 'r', encoding='utf-8') as txt_file:
            return txt_file.read().replace('\n', ' ')
    else:
        raise ValueError("Unsupported file type")

def generate_docx(summary):
    """
    Generate a DOCX file with the summary and word cloud image.

    Args:
        summary (str): The summary text to be included in the DOCX file.

    Returns:
        str: The file path of the generated DOCX file.
    """
    try:
        # Create a unique filename
        filename = f"summary_{uuid.uuid4().hex}.docx"
        file_path = os.path.join(BASE_DIR, 'downloads', filename)

        # Ensure the downloads directory exists
        os.makedirs(os.path.dirname(file_path), exist_ok=True)

        # Verify wordcloud exists
        if not os.path.exists(WORDCLOUD_PATH):
            raise FileNotFoundError(f"Wordcloud image not found at {WORDCLOUD_PATH}")

        # Create DOCX document
        doc = Document()
        doc.add_heading('Summary', level=1)
        doc.add_paragraph(summary)
        doc.add_picture(WORDCLOUD_PATH, width=docx.shared.Inches(5), height=docx.shared.Inches(6))
        doc.save(file_path)

        return file_path
    except Exception as e:
        logging.error(f"Error generating DOCX: {e}")
        raise

def safe_remove_and_render(template_name, render_template_func):
    """
    Remove generated files and render the specified template.

    Args:
        template_name (str): The name of the template to render.
        render_template_func (function): The render_template function from Flask.

    Returns:
        Response: The rendered template.
    """
    remove_files()
    return render_template_func(template_name)
