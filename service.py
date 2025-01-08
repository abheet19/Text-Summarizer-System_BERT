import re
from wordcloud import WordCloud
import matplotlib
matplotlib.use('Agg')  # Use the Agg backend to suppress the GUI warning
import matplotlib.pyplot as plt
from summarizer import Summarizer
import docx
from urllib.parse import urlparse
import requests
from newspaper import fulltext
import logging
import pathlib
import os
import PyPDF2

logging.basicConfig(level=logging.ERROR)

def clean_and_process(text):
    """
    Clean the input text and generate a word cloud image.

    Args:
        text (str): The input text to be cleaned.

    Returns:
        str: The cleaned text.
    """
    text = re.sub(r'\d', ' ', text)
    text = re.sub(r'\W', ' ', text)
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'\[[0-9]*\]', ' ', text)

    try:
        wordcloud = WordCloud(max_font_size=100, max_words=100, background_color="white").generate(text)
        plt.figure()
        plt.imshow(wordcloud, interpolation='bilinear')
        plt.axis("off")
        plt.savefig('static/img/wordcloud/wordcloud.png')
    except Exception as e:
        logging.error(f"Error generating word cloud: {e}")
    return text

def BERT(cleaned_text):
    """
    Generate a summary using the BERT model.

    Args:
        cleaned_text (str): The cleaned input text.

    Returns:
        str: The generated summary.
    """
    try:
        model = Summarizer()
        result = model(cleaned_text, min_length=30)
        return "".join(result)
    except Exception as e:
        logging.error(f"Error generating summary: {e}")
        return "Error generating summary"

def url_validator(url):
    """
    Validate the given URL.

    Args:
        url (str): The URL to be validated.

    Returns:
        bool: True if the URL is valid, False otherwise.
    """
    try:
        result = urlparse(url)
        return all([result.scheme, result.netloc, result.path])
    except Exception as e:
        logging.error(f"Error validating URL: {e}")
        return False

def dump(text):
    """
    Save the summary and word cloud image to a DOCX file.

    Args:
        text (str): The summary text to be saved.
    """
    try:
        mydoc = docx.Document()
        mydoc.add_heading("Summary", 0)
        mydoc.add_paragraph(text)
        mydoc.add_picture("static/img/wordcloud/wordcloud.png", width=docx.shared.Inches(5), height=docx.shared.Inches(6))
        mydoc.save('static/download/file.docx')
    except Exception as e:
        logging.error(f"Error saving DOCX file: {e}")

def fetch_article(url):
    """
    Fetch the article text from the given URL.

    Args:
        url (str): The URL of the article.

    Returns:
        str: The article text.
    """
    try:
        response = requests.get(url)
        return fulltext(response.text)
    except Exception as e:
        logging.error(f"Error fetching article: {e}")
        return "Error fetching article"

def summarize_url(url):
    """
    Summarize the article from the given URL.

    Args:
        url (str): The URL of the article.

    Returns:
        str: The summary of the article.
    """
    try:
        if not url_validator(url):
            return "Invalid URL"
        article_text = fetch_article(url)
        cleaned_text = clean_and_process(article_text)
        summary = BERT(cleaned_text)
        dump(summary)
        return summary
    except Exception as e:
        logging.error(f"Error summarizing URL: {e}")
        return "Error summarizing URL"

def remove_files():
    """
    Remove previous word cloud and DOCX files if they exist.
    """
    wordcloud = pathlib.Path("static/img/wordcloud/wordcloud.png")
    if wordcloud.is_file():
        os.remove(wordcloud)

    file = pathlib.Path("static/download/file.docx")
    if file.is_file():
        os.remove(file)

def process_file(file):
    """
    Process the uploaded file and return the extracted text.

    Args:
        file: The uploaded file.

    Returns:
        str: The extracted text.
    """
    suffix = pathlib.Path(file.filename).suffix
    if suffix == ".pdf":
        with open(file.filename, 'rb') as pdf:
            pdfReader = PyPDF2.PdfReader(pdf)
            return "".join([page.extract_text() for page in pdfReader.pages])
    elif suffix == ".docx":
        doc = docx.Document(file.filename)
        return "".join([para.text for para in doc.paragraphs])
    elif suffix == ".txt":
        with open(file.filename, 'r') as txt_file:
            return txt_file.read().replace('\n', '')
    else:
        raise ValueError("Unsupported file type")
