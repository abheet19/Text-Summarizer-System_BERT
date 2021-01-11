from flask import Flask, request, jsonify, render_template, redirect, url_for,send_file
from os import path
import json
from newspaper import fulltext
from wordcloud import WordCloud, STOPWORDS, ImageColorGenerator
import matplotlib.pyplot as plt
import re
from summarizer import Summarizer
import PyPDF2
import fpdf
import docx
import pathlib
from urllib.parse import urlparse
import requests
import os


app = Flask(__name__)

#--------------------helper functions for inner processing--------------------#


def clean_and_process(text):
    # cleaning the text

    clean_abstract = []
    clean_abstract = re.sub(r'\d', ' ', text)
    clean_abstract = re.sub(r'\W', ' ', clean_abstract)
    clean_abstract = re.sub(r'\s+', ' ', text)
    text = re.sub(r'\[[0-9]*\]', ' ', text)

    # saving the wordcloud for pictorial  representation

    wordcloud = WordCloud(max_font_size=100, max_words=100,
                          background_color="white").generate(str(text))
    plt.figure()
    plt.imshow(wordcloud, interpolation='bilinear')
    plt.axis("off")
    plt.savefig('static/img/wordcloud/wordcloud.png')

    return text


def BERT(cleaned_text):

    model = Summarizer()
    result = model(cleaned_text, min_length=30)
    full = "".join(result)
    return full


def url_validator(url):
    try:
        result = urlparse(url)
        return all([result.scheme, result.netloc, result.path])
    except:
        return False


def dump(text):
    mydoc = docx.Document()
    mydoc.add_heading("Summary", 0)
    mydoc.add_paragraph(text)
    mydoc.add_picture("static/img/wordcloud/wordcloud.png", width=docx.shared.Inches(5), height=docx.shared.Inches(6))
   
    mydoc.save('static/download/file.docx')






    #-----------------------------Routing methods----------------------------#


@app.route('/', methods=['GET', 'POST'])
def home():

    ################ for removing the previous wordcloud img on every request,so server dont gets loaded#############
    wordcloud = pathlib.Path("static/img/wordcloud/wordcloud.png")
    if wordcloud.is_file():
        os.remove(wordcloud)

    file = pathlib.Path("static/download/file.docx")
    if file.is_file():
        os.remove(file)

    return render_template('index.html')


@app.route('/PDF', methods=['GET', 'POST'])
def PDF():

    ################ for removing the previous wordcloud img on every request,so server dont gets loaded#############
    wordcloud = pathlib.Path("static/img/wordcloud/wordcloud.png")
    if wordcloud.is_file():
        os.remove(wordcloud)

    file = pathlib.Path("static/download/file.docx")
    if file.is_file():
        os.remove(file)


    return render_template('PDF.html')


@app.route('/PDF_result', methods=['GET', 'POST'])
def PDF_result():

    ################ for removing the previous wordcloud img on every request,so server dont gets loaded#############
    wordcloud = pathlib.Path("static/img/wordcloud/wordcloud.png")
    if wordcloud.is_file():
        os.remove(wordcloud)

    if request.method == 'POST':
        if(not request.files['name']):
            return "<h1> Please for the sake of lord , provide a file</h1>"
        else:
            f = request.files['name']
            f.save(f.filename)

            # for checking the file format
            suffix = pathlib.Path(f.filename).suffix

            # for PDF file
            if suffix == ".pdf":
                pdf = open(f.filename, 'rb')

                pdfReader = PyPDF2.PdfFileReader(pdf)
                num = pdfReader.numPages
                # indexing of pages is from (0 to num1)
                pageobj = pdfReader.pages
                text = ""
                for i in range(num):  # iterating over every page
                    # we have the raw text in the text variable
                    text += pageobj[i].extractText()
                pdf.close()
                cleaned_text = clean_and_process(text)

                # summarizer BERT model

                summary = BERT(cleaned_text)
                dump(summary)
                os.remove(f.filename)

            elif suffix == ".docx":  # for WORD file
                doc = docx.Document(f.filename)
                paras = doc.paragraphs
                text = ""
                for para in paras:
                    text += para.text

                cleaned_text = clean_and_process(text)
                # summarizer BERT model

                summary = BERT(cleaned_text)
                dump(summary)
                os.remove(f.filename)

            elif suffix == ".txt":  # for TXT file

                with open(f.filename, 'r') as file:
                    text = file.read().replace('\n', '')

                cleaned_text = clean_and_process(text)
                # summarizer BERT model

                summary = BERT(cleaned_text)
                dump(summary)
                os.remove(f.filename)

            else:
                return "<h1> Error! Please upload  PDF/WORD/TXT file</h1>"

    return render_template('PDF_result.html', summary=summary)


@app.route('/RAW', methods=['GET', 'POST'])
def RAW():
    ################ for removing the previous wordcloud img on every request,so server dont gets loaded#############
    wordcloud = pathlib.Path("static/img/wordcloud/wordcloud.png")
    if wordcloud.is_file():
        os.remove(wordcloud)

    file = pathlib.Path("static/download/file.docx")
    if file.is_file():
        os.remove(file)


    return render_template('RAW.html')


@app.route('/RAW_result', methods=['GET', 'POST'])
def RAW_result():
    ################ for removing the previous wordcloud img on every request,so server dont gets loaded#############
    wordcloud = pathlib.Path("static/img/wordcloud/wordcloud.png")
    if wordcloud.is_file():
        os.remove(wordcloud)

    if request.method == 'POST':

        url = request.form.get('name')  # get the URL from the user

        if url_validator(url) is True:  # url validation:
            text = fulltext(requests.get(url).text)
            cleaned_text = clean_and_process(text)
            # summarizer BERT model

            summary = BERT(cleaned_text)
            dump(summary)

        else:
            return "<h1>Error! Please upload a correct URL </h1>"

    return render_template('RAW_result.html', summary=summary)


@app.route('/download', methods=['GET', 'POST'])
def download():
    file = pathlib.Path("static/download/file.docx")
    if file.is_file():

        return send_file(file, as_attachment=True)


#-------------------Routing ends----------------------------#


if __name__ == "__main__":
    app.run(host='localhost',port=3000)
